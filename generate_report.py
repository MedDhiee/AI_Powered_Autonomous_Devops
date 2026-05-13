"""
Générateur du rapport technique — Incident Response Agent
Produit : rapport_incident_response_agent.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(table, col_index, width_cm):
    for row in table.rows:
        row.cells[col_index].width = Cm(width_cm)

def add_heading(doc, text, level, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return h

def add_body(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_code(doc, text):
    """Bloc de code monospace sur fond gris."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def add_table(doc, headers, rows, header_color="1F3864", alt_color="EBF0FA"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_color)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = alt_color if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)

    doc.add_paragraph()
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(level * 0.6 + 0.4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def add_note(doc, text, color_fill="FFF3CD", color_text="856404"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(color_text)
    return p

def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0C0C0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ─── Document ────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("RAPPORT TECHNIQUE")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle_p.add_run("Incident Response Agent")
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
hr(doc)
doc.add_paragraph()

meta = [
    ("Projet",    "DevOps Multi-Agent Orchestration System — CogniOps"),
    ("Module",    "devops_multi_agents/agents/incident_response/"),
    ("Auteur",    "Mohamed Dhia Hamam"),
    ("Date",      "30 avril 2026"),
    ("Version",   "1.0"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "Table des matières", 1)
toc_items = [
    "1.  Vue d'ensemble",
    "2.  Architecture des composants",
    "3.  Pipeline de traitement — 3 chemins de résolution",
    "    3.1  Path 1 — RAG (haute confiance)",
    "    3.2  Path 2 — LLM Reasoning (Groq)",
    "    3.3  Path 3 — Non résolu",
    "4.  Sources d'incidents",
    "5.  Base de connaissance RAG",
    "6.  Communication MCP",
    "7.  Traitement parallèle",
    "8.  SolutionApplier — comprendre son rôle",
    "9.  Le script .sh — comprendre son rôle",
    "10. Différence entre SolutionApplier et script .sh",
    "11. Configuration",
    "12. Format de sortie JSON",
    "13. Problèmes résolus et décisions techniques",
    "14. Limites actuelles et pistes d'amélioration",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.5 if item.startswith("    ") else 0)
    p.runs[0].font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. VUE D'ENSEMBLE
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "1. Vue d'ensemble", 1)
add_body(doc,
    "L'Incident Response Agent est un agent hybride MCP + ACP (HYBRID_MCP_ACP) qui automatise "
    "le diagnostic et la remédiation d'incidents DevOps. Il combine trois mécanismes complémentaires : "
    "une base de connaissance vectorielle (RAG via ChromaDB), un raisonnement LLM (Groq — openai/gpt-oss-120b), "
    "et une couche de validation règle-métier. Il constitue la dernière étape du pipeline d'agents du système."
)

doc.add_paragraph()
add_body(doc, "Position dans le pipeline :", bold=True)
add_code(doc,
    "Architecture → DevSecOps → CI/CD → Deployment → Chaos Engineering → Incident Response"
)

doc.add_paragraph()
add_body(doc, "Dépendances externes :", bold=True)
add_table(doc,
    ["Composant", "Technologie", "Rôle"],
    [
        ["Base RAG",     "ChromaDB (persistant)",          "Stockage vectoriel des incidents connus"],
        ["Transport RAG","MCP SDK (stdio subprocess)",      "Protocole de communication avec le serveur RAG"],
        ["LLM",          "Groq — openai/gpt-oss-120b",     "Raisonnement sur les incidents inconnus"],
        ["Source incidents","MongoDB (CogniOps.incidents)", "Incidents live en production"],
        ["Embeddings",   "all-MiniLM-L6-v2 (ChromaDB)",   "Vectorisation sémantique des incidents"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE DES COMPOSANTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "2. Architecture des composants", 1)
add_code(doc,
    "incident_response/\n"
    "├── agent.py            # Orchestrateur principal (IncidentResponseAgent)\n"
    "├── llm_reasoner.py     # Fallback LLM via Groq (openai/gpt-oss-120b)\n"
    "├── mcp_client.py       # Client MCP synchrone → ChromaDB via subprocess stdio\n"
    "├── solution_applier.py # Exécution diagnostics + génération scripts de fix\n"
    "└── __main__.py         # CLI argparse (python -m ...)"
)

add_table(doc,
    ["Fichier", "Classe / Rôle", "Responsabilité principale"],
    [
        ["agent.py",            "IncidentResponseAgent", "Orchestration du pipeline RAG → LLM → Validation → Apply"],
        ["llm_reasoner.py",     "LLMReasoner",           "Appel Groq, streaming, retry backoff, parsing JSON"],
        ["mcp_client.py",       "RagMcpClient",          "Client MCP synchrone (asyncio.run) vers ChromaDB"],
        ["solution_applier.py", "SolutionApplier",       "Exécution diagnostics read-only + génération script .sh"],
        ["__main__.py",         "main()",                "Point d'entrée CLI argparse"],
        ["mcp_servers/rag_knowledge_base.py", "IncidentKnowledgeBase", "Wrapper ChromaDB + 22 seeds + feedback loop"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 3. PIPELINE DE TRAITEMENT
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "3. Pipeline de traitement — 3 chemins de résolution", 1)
add_body(doc,
    "Chaque incident est traité individuellement par la méthode _process_incident(). "
    "Trois chemins de résolution sont possibles, parcourus dans l'ordre :"
)
add_code(doc,
    "Incident entrant\n"
    "      │\n"
    "      ▼\n"
    "┌─────────────┐\n"
    "│ RAG Search  │  ChromaDB cosine similarity, top_k=3\n"
    "│ via MCP     │\n"
    "└──────┬──────┘\n"
    "       │\n"
    "   sim ≥ seuil ?\n"
    "  ┌────┴────┐\n"
    "  OUI      NON\n"
    "  │         │\n"
    "  ▼         ▼\n"
    "Path 1   LLM Groq → Validation règle-métier\n"
    "                     ┌─────┴──────┐\n"
    "                  allowed      blocked\n"
    "                     │\n"
    "                  risk ?\n"
    "                 ┌──┴───┐\n"
    "                low   high/critical\n"
    "                 │         │\n"
    "              auto      pending_human_approval\n"
    "              apply\n"
    "                 │\n"
    "                 ▼\n"
    "           Feedback → RAG KB"
)

# 3.1
add_heading(doc, "3.1  Path 1 — RAG (haute confiance)", 2)
add_body(doc, "Condition : similarité cosinus ≥ seuil (défaut 0.7, configurable).", bold=False)
add_code(doc,
    "hits = self.mcp.search_similar_incidents(query, top_k=top_k)\n"
    "good_hits = [h for h in hits if h.get('similarity', 0) >= similarity_threshold]\n"
    "\n"
    "# Formule de conversion ChromaDB distance → similarité :\n"
    "similarity = round(1.0 - distance / 2.0, 4)   # distance ∈ [0, 2]"
)
add_body(doc, "Actions exécutées en cas de match :")
add_bullet(doc, "SolutionApplier.apply() → diagnostics read-only lancés automatiquement")
add_bullet(doc, "Script .sh de remédiation généré dans outputs/incident-fixes/")
add_bullet(doc, "Incident mis à jour dans MongoDB (champ rag_resolution)")
add_bullet(doc, "Feedback loop : _feedback_reinforce() renforce l'entrée RAG")
add_body(doc, "Résultat JSON : resolution_path = \"rag\"", italic=True)

# 3.2
add_heading(doc, "3.2  Path 2 — LLM Reasoning (Groq)", 2)
add_body(doc, "Condition : aucun hit RAG au-dessus du seuil de confiance.")
doc.add_paragraph()
add_body(doc, "Paramètres d'appel LLM :", bold=True)
add_table(doc,
    ["Paramètre", "Valeur", "Raison"],
    [
        ["model",                 "openai/gpt-oss-120b",  "Modèle de raisonnement avancé via Groq"],
        ["temperature",           "1",                    "Favorise la diversité des hypothèses"],
        ["max_completion_tokens", "1024",                 "Évite le dépassement de la limite TPM (8000)"],
        ["reasoning_effort",      "medium",               "Compromis vitesse/qualité"],
        ["stream",                "True",                 "Évite les timeouts sur longues réponses"],
    ]
)

add_body(doc, "Parsing de la réponse LLM (3 niveaux de robustesse) :", bold=True)
add_bullet(doc, "Niveau 1 : json.loads(text) — parsing direct")
add_bullet(doc, "Niveau 2 : Regex re.search(r\"\\{[\\s\\S]*\\}\", text) — extraction du bloc JSON")
add_bullet(doc, "Niveau 3 : Résultat dégradé structuré si les deux précédents échouent")

doc.add_paragraph()
add_body(doc, "Retry avec backoff exponentiel :", bold=True)
add_code(doc,
    "for attempt in range(1, 4):\n"
    "    try:\n"
    "        raw = self._stream_completion(prompt)\n"
    "        ...\n"
    "    except Exception as exc:\n"
    "        if any(code in str(exc) for code in ('429', '413', 'rate_limit', 'tokens')):\n"
    "            time.sleep(2 ** attempt)   # 2s → 4s → 8s"
)

add_body(doc, "Couche de validation règle-métier (_validate_fix) :", bold=True)
add_table(doc,
    ["Règle", "Critère", "Conséquence"],
    [
        ["Patterns bloqués",   "drop table, rm -rf /, dd if=, git push --force… (10 patterns)", "allowed=False → blocked_by_validation"],
        ["Risque élevé",       "risk_level ∈ {high, critical}",                                  "requires_approval=True → script généré, non exécuté"],
        ["Confiance LLM faible","confidence = low",                                              "Warning ajouté, traitement continue normalement"],
    ]
)

add_body(doc, "3 sous-cas du Path 2 :", bold=True)
add_table(doc,
    ["Cas", "Condition", "Action", "Statut JSON"],
    [
        ["CAS A", "allowed=True  ET requires_approval=False", "SolutionApplier.apply() + feedback RAG", "applied / pending_manual_fix"],
        ["CAS B", "allowed=True  ET requires_approval=True",  "Script .sh généré, non exécuté",          "pending_human_approval"],
        ["CAS C", "allowed=False (pattern destructeur)",      "Aucun script, aucune exécution",           "blocked_by_validation"],
    ]
)
add_body(doc, "Résultat JSON : resolution_path = \"llm\"", italic=True)

# 3.3
add_heading(doc, "3.3  Path 3 — Non résolu", 2)
add_body(doc,
    "Condition : LLM indisponible (GROQ_API_KEY absent) ou échec après 3 tentatives."
)
add_note(doc,
    "⚠  Décision importante : les incidents non résolus ne sont PAS injectés dans la base RAG. "
    "Raison : l'injection d'entrées avec \"Manual investigation required\" comme solution "
    "crée un problème de RAG self-poisoning — ces entrées se correspondent elles-mêmes "
    "à sim ≈ 0.97 lors des runs suivants, polluant tous les résultats.",
    color_fill="FFF3CD", color_text="856404"
)
add_body(doc, "Résultat JSON : resolution_path = \"none\"", italic=True)

# ════════════════════════════════════════════════════════════════════════════
# 4. SOURCES D'INCIDENTS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "4. Sources d'incidents", 1)
add_body(doc, "L'agent agrège trois sources en parallèle avant le traitement :")
add_code(doc, "all_incidents = mongo_incidents + mock_incidents + synthetic")

add_table(doc,
    ["Source", "Description", "Filtre / Condition"],
    [
        ["MongoDB",      "Collection CogniOps.incidents (production live)",       "status ∈ {open, investigating}, tri sévérité, limite 20"],
        ["Mock JSON",    "Fichier mock_incidents.json (test offline)",             "4 incidents INC-001 à INC-004"],
        ["Synthétiques", "Construits depuis la sortie DevSecOps de l'orchestrateur", "Findings critical/high → 1 incident agrégé"],
    ]
)

add_body(doc, "Exemple — mock_incidents.json (INC-001) :", bold=True)
add_code(doc,
    "{\n"
    '  "_id": "INC-001",\n'
    '  "status": "open",\n'
    '  "severity": "high",\n'
    '  "environment": "production",\n'
    '  "service_names": ["inventory-api"],\n'
    '  "summary": "Pod inventory-api-pod keeps restarting with OOMKilled",\n'
    '  "last_error_payload": [\n'
    '    {"stderr_tail": "OOMKilled exit code 137 - Not enough space errno=12"}\n'
    "  ]\n"
    "}"
)

# ════════════════════════════════════════════════════════════════════════════
# 5. BASE DE CONNAISSANCE RAG
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "5. Base de connaissance RAG", 1)
add_body(doc,
    "Le store vectoriel ChromaDB est persisté dans outputs/rag_store/. "
    "À la première initialisation, 22 incidents seed sont injectés automatiquement."
)

add_table(doc,
    ["Catégorie", "Entrées seed", "Source"],
    [
        ["Kubernetes",       "CrashLoopBackOff, ImagePullBackOff, OOMKilled, Pending, Probe Fail", "Base générale"],
        ["Docker",           "Build failure, Push failure",                                         "Base générale"],
        ["Infrastructure",   "Terraform apply fail, State lock",                                    "Base générale"],
        ["Deployment",       "Health check fail",                                                   "Base générale"],
        ["Security (Trivy)", "DB download fail, Image pull fail",                                   "trivy.pdf"],
        ["CI/CD",            "Runner offline, YAML syntax error, Secret missing",                   "ci_cd.pdf"],
        ["Gitleaks",         "False positive, Secret in git history",                               "gitleaks.pdf"],
        ["Security générale","Secret exposed, CVE critique",                                        "Base générale"],
        ["Network",          "Connection timeout, Service unavailable (503)",                       "Base générale"],
    ]
)

add_body(doc, "Embedding : chaque document est vectorisé sous la forme :", bold=True)
add_code(doc, 'text = title + "\\n" + description + "\\n" + solution')
add_body(doc,
    "La solution est incluse dans le vecteur pour maximiser la similarité sémantique "
    "lorsqu'un incident courant correspond à un incident déjà résolu."
)

add_body(doc, "Feedback loop :", bold=True)
add_body(doc,
    "Après résolution via LLM (CAS A ou CAS B), _feedback_store_llm() persiste "
    "le nouvel incident dans ChromaDB via add_incident_solution(). "
    "Les runs suivants retrouvent la solution directement par RAG sans re-solliciter le LLM."
)

# ════════════════════════════════════════════════════════════════════════════
# 6. COMMUNICATION MCP
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "6. Communication MCP", 1)
add_body(doc,
    "Le RagMcpClient lance le serveur RAG comme sous-processus stdio à chaque appel. "
    "Le protocole MCP (Model Context Protocol) est utilisé pour exposer la base ChromaDB "
    "comme un outil appelable par l'agent."
)
add_code(doc,
    "server_params = StdioServerParameters(\n"
    "    command=sys.executable,\n"
    "    args=[\"-m\", \"devops_multi_agents.mcp_servers.incidents_rag_server\"],\n"
    ")\n"
    "async with stdio_client(server_params) as (read, write):\n"
    "    async with ClientSession(read, write) as session:\n"
    "        await session.initialize()\n"
    "        result = await session.call_tool(\n"
    "            \"search_similar_incidents\", {\"query\": query, \"top_k\": top_k}\n"
    "        )"
)

add_body(doc, "Outils MCP exposés :", bold=True)
add_table(doc,
    ["Outil MCP", "Paramètres", "Retour"],
    [
        ["search_similar_incidents", "query: str, top_k: int", "list[dict] — hits triés par similarité"],
        ["add_incident_solution",    "incident_id, title, description, solution, category, severity", "bool — succès de l'insertion"],
    ]
)
add_note(doc,
    "ℹ  Le client MCP est synchrone (asyncio.run()) mais chaque thread du ThreadPoolExecutor "
    "crée son propre event loop — ce design est compatible avec le traitement parallèle.",
    color_fill="D1ECF1", color_text="0C5460"
)

# ════════════════════════════════════════════════════════════════════════════
# 7. TRAITEMENT PARALLÈLE
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "7. Traitement parallèle", 1)
add_body(doc,
    "Les incidents sont traités en parallèle via ThreadPoolExecutor. "
    "Les appels RAG (I/O sous-processus stdio) et LLM (I/O HTTP Groq) sont tous deux "
    "I/O-bound, ce qui rend le threading particulièrement efficace."
)
add_code(doc,
    "max_workers = int(os.getenv(\"INCIDENT_AGENT_WORKERS\", \"4\"))\n"
    "\n"
    "with ThreadPoolExecutor(max_workers=max_workers) as pool:\n"
    "    futures = {\n"
    "        pool.submit(self._process_incident, inc, ...): idx\n"
    "        for idx, inc in enumerate(all_incidents)\n"
    "    }\n"
    "    for future in as_completed(futures):\n"
    "        processed_map[futures[future]] = future.result()\n"
    "\n"
    "# Réordonnement : les futures complètent hors-ordre\n"
    "processed = [processed_map[i] for i in range(len(all_incidents))]"
)
add_table(doc,
    ["Mode", "Durée (24 incidents)", "Facteur"],
    [
        ["Séquentiel",          "259 secondes", "1×  (référence)"],
        ["Parallèle (4 workers)","39 secondes",  "6.5× plus rapide"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 8. SOLUTIONAPPLIER — COMPRENDRE SON RÔLE
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "8. SolutionApplier — comprendre son rôle", 1)
add_body(doc,
    "Le SolutionApplier est la classe Python chargée d'agir concrètement sur le système "
    "une fois qu'une solution a été trouvée (par RAG ou LLM). "
    "Il distingue rigoureusement deux types d'actions."
)

add_body(doc, "Deux types d'actions :", bold=True)
add_table(doc,
    ["Type", "Nature", "Exemples", "Comportement"],
    [
        ["Diagnostics", "Lecture seule — aucun risque",     "kubectl logs, kubectl describe, docker images, terraform plan", "Exécutés automatiquement sans confirmation"],
        ["Fix commands", "Modification du système — risqué", "kubectl apply -f, kubectl rollout restart, terraform apply",   "Écrits dans le script .sh, jamais exécutés sans --auto-apply"],
    ]
)

add_body(doc, "Commandes considérées comme sûres (auto-exécutées) :", bold=True)
add_code(doc,
    "_SAFE_PREFIXES = (\n"
    '    "kubectl logs",\n'
    '    "kubectl describe",\n'
    '    "kubectl get",\n'
    '    "kubectl rollout status",\n'
    '    "kubectl top",\n'
    '    "docker images",\n'
    '    "docker system df",\n'
    '    "terraform show",\n'
    '    "terraform plan",\n'
    ")"
)

add_body(doc, "Interpolation des templates de commandes :", bold=True)
add_body(doc,
    "Les commandes du catalogue RAG contiennent des placeholders ({service_name}, {pod_name}, etc.) "
    "qui sont remplacés par les valeurs réelles de l'incident :"
)
add_code(doc,
    "ctx = {\n"
    '    "service_name": "inventory-api",\n'
    '    "pod_name":     "inventory-api-pod",\n'
    '    "manifest_path":"k8s/inventory-api.yaml",\n'
    '    "environment":  "production",\n'
    "}\n"
    "cmd = \"kubectl apply -f {manifest_path}\"\n"
    "→    \"kubectl apply -f k8s/inventory-api.yaml\""
)

add_body(doc, "Flux d'exécution interne de apply() :", bold=True)
add_code(doc,
    "SolutionApplier.apply(incident, rag_match)\n"
    "  │\n"
    "  ├── 1. Construire ctx (interpolation)\n"
    "  ├── 2. Pour chaque auto_command → si _is_safe() → subprocess.run()\n"
    "  ├── 3. Pour chaque fix_command  → si auto_apply → subprocess.run()\n"
    "  │                                 sinon        → pending_fix_commands[]\n"
    "  ├── 4. _write_fix_script() → fichier .sh sur disque\n"
    "  └── 5. Retourner rapport structuré (dict)"
)

# ════════════════════════════════════════════════════════════════════════════
# 9. LE SCRIPT .SH — COMPRENDRE SON RÔLE
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "9. Le script .sh — comprendre son rôle", 1)
add_body(doc,
    "Le script .sh est un fichier bash généré automatiquement dans outputs/incident-fixes/. "
    "Il sert de document de travail entre l'agent et l'opérateur humain. "
    "Il n'est jamais exécuté automatiquement par défaut."
)

add_body(doc, "Structure d'un script généré (exemple INC-001 — OOMKilled) :", bold=True)
add_code(doc,
    "#!/usr/bin/env bash\n"
    "# ============================================================\n"
    "# Auto-generated incident fix script\n"
    "# Generated : 2026-04-29T10:32:00Z\n"
    "# Incident  : INC-001\n"
    "# RAG match : OOMKilled – out of memory  (similarity=0.905)\n"
    "# Category  : kubernetes   Severity: high\n"
    "# ============================================================\n"
    "\n"
    "# ── Incident details ──────────────────────────────────────────\n"
    "# Status    : open\n"
    "# Services  : inventory-api\n"
    "# Env       : production\n"
    "# Occurred  : 7x\n"
    "\n"
    "# ── Remediation plan ──────────────────────────────────────────\n"
    "# 1. Check current limits: kubectl describe pod <pod> | grep -A3 Limits\n"
    "# 2. Increase memory limit in deployment manifest\n"
    "# 3. Profile the app for memory leaks\n"
    "# 4. Apply updated manifest: kubectl apply -f deployment.yaml\n"
    "# 5. Monitor: kubectl top pod <pod>\n"
    "\n"
    "# ── Diagnostic commands (already executed) ────────────────────\n"
    "# CMD: kubectl describe pod inventory-api-pod\n"
    "# RC : 0\n"
    "#   Limits:\n"
    "#     memory: 256Mi    ← trop peu pour la JVM\n"
    "#   Last State: OOMKilled\n"
    "\n"
    "# ── Fix commands (review before running!) ─────────────────────\n"
    "set -euo pipefail\n"
    "\n"
    "kubectl apply -f k8s/inventory-api.yaml\n"
    "\n"
    "echo 'Fix script completed.'"
)

add_body(doc, "Sections du script :", bold=True)
add_table(doc,
    ["Section", "Contenu", "Généré par"],
    [
        ["En-tête",          "Timestamp, incident ID, RAG match + similarity score, sévérité",    "_write_fix_script()"],
        ["Incident details",  "Status, services, environnement, nombre d'occurrences",             "_write_fix_script()"],
        ["Remediation plan",  "Étapes commentées extraites de la solution RAG/LLM",               "_write_fix_script()"],
        ["Diagnostic output", "Sorties des commandes read-only déjà exécutées (tronquées 4000c)", "SolutionApplier.apply()"],
        ["Fix commands",      "Commandes de remédiation à réviser et exécuter manuellement",      "_write_fix_script()"],
    ]
)

add_note(doc,
    "💡  Nommage du fichier : fix_<incident_id>_<timestamp_UTC>.sh\n"
    "    Exemple : fix_INC-001_20260429T103200Z.sh\n"
    "    Emplacement : devops_multi_agents/outputs/incident-fixes/",
    color_fill="D4EDDA", color_text="155724"
)

# ════════════════════════════════════════════════════════════════════════════
# 10. DIFFÉRENCE ENTRE SOLUTIONAPPLIER ET SCRIPT .SH
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "10. Différence entre SolutionApplier et script .sh", 1)
add_body(doc,
    "La confusion vient du fait que les deux semblent \"appliquer la solution\". "
    "En réalité ils ont des natures et des cycles de vie totalement différents."
)

add_table(doc,
    ["Critère", "SolutionApplier", "Script .sh"],
    [
        ["Nature",       "Classe Python (code actif)",           "Fichier bash (document statique)"],
        ["Quand ?",      "Pendant l'exécution de l'agent",       "Après l'exécution, relu par l'humain"],
        ["Qui l'utilise?","L'agent Python automatiquement",       "L'opérateur humain manuellement"],
        ["Ce qu'il fait","Orchestre diagnostics + génère le .sh", "Fournit les commandes à exécuter"],
        ["Exécution",    "Lance subprocess.run() pour diagnostics","Exécuté manuellement via bash"],
        ["Auto-apply ?", "Oui (diagnostics) / optionnel (fixes)", "Non — toujours manuel"],
        ["Durée de vie", "Objet Python en RAM pendant le run",    "Fichier persistant sur disque"],
        ["Destinataire", "Interne à l'agent",                     "L'opérateur / équipe on-call"],
    ]
)

add_body(doc, "Schéma de relation :", bold=True)
add_code(doc,
    "┌─────────────────────────────────────────────────────────┐\n"
    "│                    SolutionApplier                      │\n"
    "│                   (classe Python)                       │\n"
    "│                                                         │\n"
    "│  1. Lance les diagnostics → résultats en RAM            │\n"
    "│  2. Génère le script .sh  → fichier sur disque          │\n"
    "│  3. Optionnellement lance les fix commands              │\n"
    "└─────────────────────────────────────────────────────────┘\n"
    "                          │ génère\n"
    "                          ▼\n"
    "┌─────────────────────────────────────────────────────────┐\n"
    "│            Script fix_INC-001_20260429.sh               │\n"
    "│                  (fichier bash statique)                │\n"
    "│                                                         │\n"
    "│  • Résumé de l'incident                                 │\n"
    "│  • Plan de remédiation commenté                         │\n"
    "│  • Sortie des diagnostics déjà exécutés                 │\n"
    "│  • Commandes à exécuter manuellement                    │\n"
    "└─────────────────────────────────────────────────────────┘\n"
    "                          │ lu par\n"
    "                          ▼\n"
    "                    Opérateur humain\n"
    "                  (modifie, valide, exécute)"
)

add_note(doc,
    "En résumé : le SolutionApplier est le robot qui agit et prépare le travail. "
    "Le script .sh est le bon de travail qu'il laisse sur le bureau. "
    "La séparation existe car observer est toujours sûr, "
    "mais modifier la production nécessite un œil humain.",
    color_fill="E2D9F3", color_text="4A0E8F"
)

# ════════════════════════════════════════════════════════════════════════════
# 11. CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "11. Configuration", 1)

add_table(doc,
    ["Variable d'environnement", "Défaut", "Description"],
    [
        ["GROQ_API_KEY",              "—",                          "Clé API Groq (obligatoire pour le LLM)"],
        ["RAG_SIMILARITY_THRESHOLD",  "0.7",                        "Seuil minimum de similarité cosinus RAG"],
        ["RAG_PERSIST_DIR",           "outputs/rag_store/",         "Répertoire de persistance ChromaDB"],
        ["INCIDENT_AGENT_WORKERS",    "4",                          "Nombre de threads parallèles"],
        ["MONGO_URI",                 "mongodb://localhost:27017/",  "URI de connexion MongoDB"],
        ["MONGO_DB_NAME",             "CogniOps",                   "Nom de la base de données MongoDB"],
    ]
)

add_body(doc, "Commande CLI complète :", bold=True)
add_code(doc,
    "python -m devops_multi_agents.agents.incident_response \\\n"
    "  --arch    outputs/architecture-analysis.json  \\\n"
    "  --security outputs/devsecops-security.json    \\\n"
    "  --chaos   outputs/chaos-experiments.json      \\\n"
    "  --mock-incidents devops_multi_agents/mock_incidents.json \\\n"
    "  --similarity-threshold 0.7  \\\n"
    "  --top-k 3                   \\\n"
    "  --auto-apply                \\\n"
    "  --output outputs/incident-response.json"
)

# ════════════════════════════════════════════════════════════════════════════
# 12. FORMAT DE SORTIE JSON
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "12. Format de sortie JSON", 1)
add_code(doc,
    "{\n"
    '  "protocol": "hybrid_mcp_acp",\n'
    '  "summary": "Processed 24 incident(s): 24 via RAG, 0 via LLM, 0 unresolved.",\n'
    '  "total_incidents": 24,\n'
    '  "rag_matched": 24,\n'
    '  "llm_resolved": 0,\n'
    '  "unresolved": 0,\n'
    '  "similarity_threshold": 0.7,\n'
    '  "processed_incidents": [\n'
    "    {\n"
    '      "incident_id": "INC-001",\n'
    '      "resolution_path": "rag",\n'
    '      "severity": "high",\n'
    '      "services": ["inventory-api"],\n'
    '      "rag_hits": [{"title": "OOMKilled", "similarity": 0.905}],\n'
    '      "apply_result": {\n'
    '        "fix_script": "outputs/incident-fixes/fix_INC-001_*.sh",\n'
    '        "status": "pending_manual_fix",\n'
    '        "pending_fix_commands": ["kubectl apply -f k8s/inventory-api.yaml"]\n'
    "      }\n"
    "    }\n"
    "  ],\n"
    '  "runbook": [...],\n'
    '  "incident_actions": [...]\n'
    "}"
)

# ════════════════════════════════════════════════════════════════════════════
# 13. PROBLÈMES RÉSOLUS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "13. Problèmes résolus et décisions techniques", 1)

add_table(doc,
    ["Problème rencontré", "Solution implémentée"],
    [
        ["RAG self-poisoning (sim ≈ 0.97 sur incidents non résolus)",
         "Suppression de _ingest_unmatched_incident() du Path 3"],
        ["GROQ_BASE_URL doublé → 404 Not Found",
         "base_url='https://api.groq.com' hardcodé dans le constructeur Groq (SDK ajoute /openai/v1 lui-même)"],
        ["Dépassement TPM — 413 Token limit (8694 > 8000 tokens)",
         "Prompt système raccourci (<200 tokens), max_completion_tokens=1024"],
        ["LLM retourne JSON imbriqué dans une string",
         "Double parsing : direct → regex {[\\s\\S]*} → résultat dégradé structuré"],
        ["Path LLM sans script réel (faux statut 'llm_fix_applied')",
         "Résultats LLM routés via SolutionApplier avec un llm_as_rag synthétique"],
        ["Traitement séquentiel trop lent (259s pour 24 incidents)",
         "ThreadPoolExecutor(max_workers=4) → 39s (gain 6.5×)"],
        ["Pas de retry sur rate-limit Groq (429/413)",
         "Backoff exponentiel : 3 tentatives avec sleep(2^attempt)"],
        ["INC-003/INC-004 matchaient de mauvaises entrées RAG",
         "+7 seeds extraits des PDFs de connaissance (trivy.pdf, ci_cd.pdf, gitleaks.pdf)"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# 14. LIMITES ET AMÉLIORATIONS
# ════════════════════════════════════════════════════════════════════════════

add_heading(doc, "14. Limites actuelles et pistes d'amélioration", 1)

add_table(doc,
    ["Limite actuelle", "Amélioration suggérée"],
    [
        ["asyncio.run() dans chaque thread → sous-processus MCP par appel",
         "Pool de workers async ou client MCP persistant partagé entre threads"],
        ["Feedback loop _feedback_reinforce est un simple log",
         "Implémenter un compteur de hits MongoDB pour prioriser les solutions fréquentes"],
        ["Seuil de similarité unique pour tous les types d'incidents",
         "Seuils différenciés par catégorie (kubernetes: 0.80, security: 0.75, network: 0.70)"],
        ["Confiance LLM non calibrée sur données historiques",
         "Calibration sur un dataset d'incidents résolus pour valider les scores"],
        ["Pas de vérification post-fix automatique",
         "Intégrer une étape de health check après application pour confirmer la résolution"],
        ["Pas d'interface de supervision en temps réel",
         "Dashboard Frontend (déjà en cours) avec popup incident + flow RAG/LLM interactif"],
    ]
)

# ════════════════════════════════════════════════════════════════════════════
# PIED DE PAGE
# ════════════════════════════════════════════════════════════════════════════

hr(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Rapport généré le 30 avril 2026  —  Mohamed Dhia Hamam  —  Projet PFE TALAN")
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

# ════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ════════════════════════════════════════════════════════════════════════════

output_path = "devops_multi_agents/outputs/rapport_incident_response_agent.docx"
doc.save(output_path)
print(f"Rapport généré : {output_path}")
